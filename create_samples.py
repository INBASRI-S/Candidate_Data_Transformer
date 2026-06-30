import os
import json
import pandas as pd
import docx

def generate_samples():
    os.makedirs("samples", exist_ok=True)
    print("Generating mock files in 'samples/' directory...")

    # 1. Generate recruiter.csv
    csv_data = [{
        "Full Name": "Alex Mercer",
        "Email": "alex.mercer@gmail.com",
        "Phone": "+1 (555) 019-2834",
        "Skills": "Python, SQL, HTML, CSS",
        "Country": "Canada",
        "Institution": "University of Toronto",
        "Degree": "Bachelor of Science",
        "Company": "Innovate LLC",
        "Title": "Intern Developer",
        "Description": "Assisted in building Python backend applications."
    }]
    pd.DataFrame(csv_data).to_csv("samples/recruiter.csv", index=False)
    print("- Created samples/recruiter.csv")

    # 2. Generate ats.json
    ats_data = {
        "Name": "Alex Mercer",
        "Emails": ["alex.mercer@ats-profile.org"],
        "Phones": ["+1-555-019-2834"],
        "Skills": ["Python", "SQL", "Docker", "Git"],
        "Location": "Canada",
        "Education": [
            {
                "Institution": "University of Toronto",
                "Degree": "B.S. in Computer Science",
                "Field of Study": "Computer Science",
                "Start Date": "2018",
                "End Date": "2022"
            }
        ],
        "Experience": [
            {
                "Company": "Innovate LLC",
                "Title": "Junior software developer",
                "Start Date": "2022-05",
                "End Date": "2023-08",
                "Description": "Developed REST APIs in Python and SQL."
            }
        ]
    }
    with open("samples/ats.json", "w") as f:
        json.dump(ats_data, f, indent=2)
    print("- Created samples/ats.json")

    # 3. Generate resume.docx
    doc = docx.Document()
    doc.add_heading("Alex Mercer", level=0)
    doc.add_paragraph("alex.mercer@gmail.com | +1 555-019-2834 | Toronto, Canada")
    
    doc.add_heading("Summary", level=1)
    doc.add_paragraph("Highly motivated developer with experience in Python backend development and cloud technologies.")
    
    doc.add_heading("Skills", level=1)
    doc.add_paragraph("Python, JavaScript, TypeScript, FastAPI, PostgreSQL, Docker, AWS, Git")
    
    doc.add_heading("Education", level=1)
    doc.add_paragraph("University of Toronto\nBachelor of Science in Computer Science, 2018 - 2022")
    
    doc.add_heading("Experience", level=1)
    doc.add_paragraph("Innovate LLC\nBackend Software Engineer, 2022 - 2024\nDeveloped microservices using FastAPI, Python, and PostgreSQL. Deployed containers on AWS using Docker.")
    
    doc.save("samples/resume.docx")
    print("- Created samples/resume.docx")

    # 4. Generate notes.txt
    notes_content = """
    Candidate name: Alex Mercer
    Contact: alex.mercer@gmail.com
    Phone: +1 555-019-2834
    Skills: Python, FastAPI, Docker, PostgreSQL, Kubernetes, Agile
    Location: Canada
    Recruiter feedback: Excellent communication skills, strong backend engineering foundation. Knows Kubernetes.
    """
    with open("samples/notes.txt", "w", encoding="utf-8") as f:
        f.write(notes_content.strip())
    print("- Created samples/notes.txt")

    # 5. Generate projection_config.json
    config_data = {
        "select_fields": ["full_name", "emails", "phones", "skills", "experience", "education", "country"],
        "rename_fields": {
            "full_name": "candidate_name",
            "emails": "contact_emails",
            "phones": "contact_numbers",
            "skills": "technical_skills"
        },
        "include_confidence": True,
        "missing_field_behavior": "default",
        "default_values": {
            "country": "Canada",
            "full_name": "Alex Mercer"
        }
    }
    with open("samples/projection_config.json", "w") as f:
        json.dump(config_data, f, indent=2)
    print("- Created samples/projection_config.json")

    print("\nAll mock files generated in 'samples/'. You can upload them using the web UI.")

if __name__ == "__main__":
    generate_samples()
